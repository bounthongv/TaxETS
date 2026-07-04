#!/usr/bin/env python3
"""Convert user-manual.md to DOCX with placeholder images."""
import re
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Saysettha OT'  # Lao-friendly font
font.size = Pt(11)

# Helper functions
def add_heading(text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    return h

def add_placeholder(title):
    """Add a screenshot placeholder box."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'[ {title} ]')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.font.italic = True
    # Add a border simulation
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {
        qn('w:top'): qn('w:val') + '="single" ' + qn('w:sz') + '="4" ' + qn('w:space') + '="1" ' + qn('w:color') + '="CCCCCC"',
    })
    pPr.append(pBdr)
    return p

def add_table(headers, rows):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(9)
    # Rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    return table

def add_code(text):
    """Add a code/monospace paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p

# ============================================================
# DOCUMENT CONTENT
# ============================================================

# Title
title = doc.add_heading('ຄູ່ມືຜູ້ໃຊ້ Tax-ETS', level=0)
for run in title.runs:
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

# Section 1
add_heading('1. ຈຸດປະສົງຂອງຄູ່ມື', 1)
doc.add_paragraph(
    'ຄູ່ມືນີ້ອະທິບາຍວິທີໃຊ້ງານລະບົບ Tax-ETS ສໍາລັບຜູ້ໃຊ້ທົ່ວໄປ. '
    'ເນື້ອຫາຈະເນັ້ນໃສ່ຂັ້ນຕອນການເຮັດວຽກຫຼັກ ເຊັ່ນ ການນໍາເຂົ້າຂໍ້ມູນ, '
    'ການກວດສອບ Batch, ການຄໍານວນ TE, ແລະການເບິ່ງລາຍງານ.'
)
doc.add_paragraph(
    'ແນວທາງຂອງຄູ່ມືນີ້ຈະບໍ່ອະທິບາຍຊໍ້າກັນທຸກປະເພດພາສີ, '
    'ເພາະຂັ້ນຕອນພື້ນຖານຄ້າຍຄືກັນ. '
    'ຈະອະທິບາຍລາຍລະອຽດຄັ້ງດຽວໃນຮູບແບບ Workflow, '
    'ແລ້ວສະຫຼຸບຈຸດທີ່ແຕກຕ່າງຂອງແຕ່ລະປະເພດພາສີ.'
)

# Section 2
add_heading('2. ຜູ້ໃຊ້ທີ່ເໝາະສົມ', 1)
doc.add_paragraph('ຄູ່ມືນີ້ເໝາະສໍາລັບ:')
items = [
    'ຜູ້ນໍາເຂົ້າຂໍ້ມູນຈາກ Excel',
    'ຜູ້ກວດສອບຂໍ້ມູນ Batch',
    'ຜູ້ດໍາເນີນການຄໍານວນ Tax Expenditure',
    'ຜູ້ເບິ່ງ ແລະ Export ລາຍງານ',
    'ຜູ້ປະສານງານກັບຜູ້ຊ່ຽວຊານດ້ານພາສີ',
]
for item in items:
    doc.add_paragraph(f'• {item}')

# Section 3
add_heading('3. ຄໍາສັບສໍາຄັນ', 1)
add_table(
    ['ຄໍາສັບ', 'ຄວາມໝາຍ'],
    [
        ['TE', 'Tax Expenditure, ມູນຄ່າລາຍຮັບທີ່ລັດສູນເສຍຈາກມາດຕະການພາສີ'],
        ['Batch', 'ກຸ່ມຂໍ້ມູນທີ່ນໍາເຂົ້າ ຫຼື ປ້ອນດ້ວຍມືພ້ອມກັນ'],
        ['Import', 'ການນໍາເຂົ້າຂໍ້ມູນຈາກ Excel'],
        ['Template', 'ແບບຟອມ Excel ສໍາລັບນໍາເຂົ້າຂໍ້ມູນ'],
        ['System TE', 'ຄ່າ TE ທີ່ລະບົບ Tax-ETS ຄໍານວນ'],
        ['Benchmark', 'ອັດຕາຫຼືຫຼັກການມາດຕະຖານທີ່ໃຊ້ເພື່ອຄໍານວນ'],
        ['Provision', 'ມາດຕາ ຫຼື ຂໍ້ກໍານົດທາງກົດໝາຍທີ່ກ່ຽວກັບ TE'],
    ]
)

# Section 4
add_heading('4. ການເຂົ້າລະບົບ', 1)

p = doc.add_paragraph()
run = p.add_run('System URL: ')
run.bold = True
run = p.add_run('https://tax-ets.apis.com.la/')
run.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)

steps = [
    'ເປີດ Browser ແລ້ວເຂົ້າ URL: https://tax-ets.apis.com.la/',
    'ໃສ່ Email ແລະ Password.',
    'ກົດປຸ່ມ Login.',
    'ຫຼັງຈາກເຂົ້າລະບົບສໍາເລັດ ລະບົບຈະເຂົ້າສູ່ Dashboard.',
]
for i, step in enumerate(steps, 1):
    doc.add_paragraph(f'{i}. {step}')

# Add user accounts table
p = doc.add_paragraph()
run = p.add_run('ບັນຊີຜູ້ໃຊ້ສໍາລັບການທົດສອບ:')
run.bold = True
add_table(
    ['Email', 'Password', 'ສິດ'],
    [
        ['trainer@example.com', 'trainer123', 'ADMIN (ສໍາລັບຜູ້ທົດສອບ ແລະ ຝຶກອົບຮົມ)'],
    ]
)

add_placeholder('Screenshot: Login page')

# Section 5
add_heading('5. ໂຄງສ້າງເມນູຫຼັກ', 1)
doc.add_paragraph('ຫຼັງຈາກ Login, ຜູ້ໃຊ້ຈະເຫັນ Sidebar ດ້ານຊ້າຍ.')
add_table(
    ['ເມນູ', 'ຈຸດປະສົງ'],
    [
        ['Dashboard', 'ໜ້າສະຫຼຸບພາບລວມ'],
        ['Get Tax Data by Import from Excel', 'ນໍາເຂົ້າຂໍ້ມູນຈາກ Excel'],
        ['Batch Management', 'ຈັດການ Batch ທີ່ນໍາເຂົ້າແລ້ວ'],
        ['TE Calculation', 'ຄໍານວນ Tax Expenditure'],
        ['TE Reports', 'ເບິ່ງ ແລະ Export ລາຍງານ'],
    ]
)
add_placeholder('Screenshot: Main sidebar navigation')

# Section 6
add_heading('6. ຂັ້ນຕອນການເຮັດວຽກຫຼັກ', 1)
doc.add_paragraph('ການໃຊ້ງານລະບົບສ່ວນໃຫຍ່ຈະເປັນຂັ້ນຕອນດັ່ງນີ້:')
workflow = [
    'Download Template.',
    'ຕື່ມຂໍ້ມູນໃນ Excel Template.',
    'Import Excel ເຂົ້າລະບົບ.',
    'ກວດສອບຜົນການ Import ແລະ Download Log ຖ້າມີ Error.',
    'ກົດ Go to TE Calculation ເພື່ອເຂົ້າໜ້າຄໍານວນ.',
    'ຢູ່ໜ້າ TE Calculation, ກົດ Run TE Calculation.',
    'ກວດສອບຜົນການຄໍານວນ.',
    'ເບິ່ງລາຍງານ ແລະ Export Excel/PDF.',
]
for i, step in enumerate(workflow, 1):
    doc.add_paragraph(f'{i}. {step}')

add_code('Excel Template → Import → View Batch → Go to TE Calculation → Run TE Calculation → Reports')

note = doc.add_paragraph()
run = note.add_run('ຫມາຍສໍາຄັນ: ')
run.bold = True
run = note.add_run(
    'ປຸ່ມ "Go to TE Calculation" ຢູ່ໜ້າ View ແມ່ນສໍາລັບນໍາທາງໄປໜ້າຄໍານວນ. '
    'ປຸ່ມ "Run TE Calculation" ຢູ່ໜ້າ TE Calculation ແມ່ນສັ່ງໃຫ້ຄໍານວນຕົວຈິງ.'
)

# Section 7 - Detailed Workflow
add_heading('7. ຕົວຢ່າງ Workflow ລະອຽດ: Domestic VAT', 1)
doc.add_paragraph('ພາກນີ້ໃຊ້ Domestic VAT ເປັນຕົວຢ່າງລະອຽດ. ສໍາລັບປະເພດພາສີອື່ນ ຂັ້ນຕອນຈະຄ້າຍຄືກັນ.')

# 7.1
add_heading('7.1 Download Template', 2)
add_code('Get Tax Data by Import from Excel\n  > Data Requirement to estimate TE\n    > Domestic VAT')
doc.add_paragraph('• ກົດ Download Template ແລະ ບັນທຶກ Excel Template ໄວ້ໃນເຄື່ອງ.')
add_placeholder('Screenshot: Domestic VAT import page and Download Template button')

# 7.2
add_heading('7.2 ການກຽມ Excel', 2)
doc.add_paragraph('ໃຫ້ຕື່ມຂໍ້ມູນໃນ Template ຕາມຮູບແບບທີ່ກໍານົດ.')
doc.add_paragraph('ຂໍ້ຄວນລະວັງ:')
tips = [
    'ຢ່າປ່ຽນຊື່ຄໍລໍາ.',
    'ຢ່າລຶບ Sheet ຫຼື ຄໍລໍາທີ່ Template ກໍານົດ.',
    'ກວດສອບ Province/District ໃຫ້ຖືກຕາມ Dictionary.',
    'ວັນທີຄວນເປັນ Date ທີ່ Excel ອ່ານໄດ້.',
    'ຈໍານວນເງິນຄວນເປັນຕົວເລກ, ບໍ່ແມ່ນຂໍ້ຄວາມ.',
]
for tip in tips:
    doc.add_paragraph(f'• {tip}')

# 7.3
add_heading('7.3 Import Excel', 2)
doc.add_paragraph('1. ຢູ່ໜ້າ Domestic VAT Import, ເລືອກໄຟລ໌ Excel.')
doc.add_paragraph('2. ກົດ Import.')
doc.add_paragraph('3. ລະບົບຈະສ້າງ Batch ID ໃໝ່.')
doc.add_paragraph('4. ຖ້າ Import ສໍາເລັດ ຈະເຫັນ Batch ໃນລາຍການ Recent Batches.')
doc.add_paragraph('5. ຖ້າມີ Error, ກົດ Download Log ເພື່ອກວດສອບ.')
add_placeholder('Screenshot: Import file selection and recent batches')

# 7.4
add_heading('7.4 ເບິ່ງ Batch ໃນ Batch Management', 2)
add_code('Get Tax Data by Import from Excel\n  > Data Requirement to estimate TE\n    > Batch Management')
doc.add_paragraph('ໃຊ້ Filter Tax Type ເລືອກ Domestic VAT, ຄົ້ນຫາ Batch ID, ກົດ View.')
add_placeholder('Screenshot: Batch Management page')

# 7.5
add_heading('7.5 ເບິ່ງ ແລະ ແກ້ໄຂຂໍ້ມູນໃນ Batch', 2)
doc.add_paragraph('• ກວດສອບຈໍານວນ Records.')
doc.add_paragraph('• ໃຊ້ Search, Filter, Edit ຕາມຄວາມຕ້ອງການ.')
doc.add_paragraph('• ກົດ Add Record to Batch ຖ້າຕ້ອງການເພີ່ມລາຍການ.')
add_placeholder('Screenshot: Domestic VAT records page')

# 7.6
add_heading('7.6 Go to TE Calculation', 2)
doc.add_paragraph('ຢູ່ໜ້າ Domestic VAT Records, ກົດ Go to TE Calculation (ປຸ່ມສີຂຽວ).')
add_placeholder('Screenshot: TE Calculation page - initial state')

# 7.7
add_heading('7.7 Run TE Calculation', 2)
doc.add_paragraph('1. ກົດ Run TE Calculation.')
doc.add_paragraph('2. ລໍຖ້າລະບົບຄໍານວນ (ຈະມີ Spinner).')
doc.add_paragraph('3. ກວດສອບຜົນສະຫຼຸບ.')
add_placeholder('Screenshot: TE Calculation results')

# Section 8
add_heading('8. ການໃຊ້ Batch Management', 1)
add_code('Get Tax Data by Import from Excel\n  > Data Requirement to estimate TE\n    > Batch Management')
doc.add_paragraph('ໃນໜ້ານີ້ຜູ້ໃຊ້ສາມາດ:')
bm_items = [
    'ເບິ່ງ Batch ທັງໝົດ, Filter ຕາມ Tax Type, Search Batch ID',
    'ເບິ່ງຈໍານວນ Rows, Years, Import Date',
    'ເປີດ View, TE Calculation, Download Log, Delete Batch',
]
for item in bm_items:
    doc.add_paragraph(f'• {item}')
add_placeholder('Screenshot: Batch Management page')
doc.add_paragraph('ຫມາຍ: ການ Delete Batch ຈະລຶບຂໍ້ມູນທັງໝົດຂອງ Batch ນັ້ນ. ຢືນຢັນກ່ອນລຶບສະເໝີ.')

# Section 9
add_heading('9. ຂໍ້ມູນເຉພາະຂອງແຕ່ລະປະເພດພາສີ', 1)
doc.add_paragraph('ຕາຕະລາງຕໍ່ໄປນີ້ສະຫຼຸບຈຸດສໍາຄັນຂອງແຕ່ລະປະເພດພາສີ.')
add_table(
    ['ປະເພດ', 'Template Cols', 'ຈຸດສັງເກດ'],
    [
        ['Profit Tax', 'Custom', 'ໃຊ້ຂໍ້ມູນບໍລິສັດ, ກໍາໄລ, ອັດຕາ Benchmark'],
        ['PIT', '35 cols', 'ອີງໃສ່ PTIN/TIN ແລະ Tax Year'],
        ['Domestic VAT', '17 cols', 'ໃຊ້ຂໍ້ມູນ Filing Period, Sales, VAT'],
        ['SEZ Developer', '20 cols', 'ໃຊ້ template ດຽວກັບ Investor'],
        ['SEZ Investor', '20 cols', 'ກວດ Column J = Yes'],
        ['Land Concession', '19 cols', 'Cascading District dropdown'],
        ['Resource Fee', '17 cols', 'ອີງໃສ່ປະເພດຊັບພະຍາກອນ'],
        ['Royalty Fee', '16 cols', 'ອີງໃສ່ມູນຄ່າຂາຍໄຟຟ້າ'],
        ['ASYCUDA', '46 cols', 'Import ຄັ້ງດຽວແຍກ 3 ປະເພດ TE'],
    ]
)

# Section 10
add_heading('10. ການ Run TE Calculation', 1)
doc.add_paragraph('ມີ 2 ຂັ້ນຕອນ:')
p = doc.add_paragraph()
run = p.add_run('ຂັ້ນຕອນທີ 1: Go to TE Calculation. ')
run.bold = True
p.add_run('ຈາກໜ້າ View Batch ກົດ Go to TE Calculation ຫຼື ຈາກເມນູ TE Calculation.')

p = doc.add_paragraph()
run = p.add_run('ຂັ້ນຕອນທີ 2: Run TE Calculation. ')
run.bold = True
p.add_run('ກົດ Run TE Calculation ຢູ່ໜ້າ TE Calculation. ລໍຖ້າ Processing... ແລະ ກວດສອບຜົນ.')

# Section 11
add_heading('11. ການໃຊ້ລາຍງານ', 1)
add_code('TE Reports')
add_table(
    ['ລາຍງານ', 'ຈຸດປະສົງ'],
    [
        ['TE by Tax Type', 'ສະຫຼຸບ TE ຕາມປະເພດພາສີ'],
        ['TE by Sector', 'ສະຫຼຸບ TE ຕາມຂະແໜງ'],
        ['TE by Location', 'ສະຫຼຸບ TE ຕາມແຂວງ'],
        ['TE (% of GDP)', 'ປຽບທຽບ TE ກັບ GDP'],
        ['TE (% of Revenue)', 'ປຽບທຽບ TE ກັບ Revenue'],
        ['TE by Provision', 'ສະຫຼຸບ TE ຕາມ Provision'],
        ['Customs Regime', 'ສະຫຼຸບ TE ຂອງ ASYCUDA'],
    ]
)

# Sections 12-18 (simplified)
for sec_num, title_text in [
    ('12', 'ຂໍ້ຜິດພາດທີ່ພົບເລື້ອຍ'),
    ('13', 'ການໃຊ້ Download Log'),
    ('14', 'ຂໍ້ຄວນລະວັງກ່ອນລຶບຂໍ້ມູນ'),
    ('15', 'ຄໍາແນະນໍາການເຮັດວຽກທີ່ດີ'),
    ('16', 'ສິ່ງທີ່ຄວນແຈ້ງ Admin ຫຼື Developer'),
]:
    add_heading(f'{sec_num}. {title_text}', 1)

# Error table
add_table(
    ['ອາການ', 'ສາເຫດ', 'ວິທີແກ້'],
    [
        ['Import ບໍ່ສໍາເລັດ', 'Template ຜິດ / ຂໍ້ມູນຂາດ', 'Download Log ແລະກວດ Excel'],
        ['Province Error', 'ຊື່ແຂວງບໍ່ກົງ', 'ແກ້ໃນ Excel ຫຼື ແຈ້ງ Admin'],
        ['Date Error', 'Format ວັນທີບໍ່ຖືກ', 'ແກ້ໃຫ້ເປັນ Date ໃນ Excel'],
        ['ລາຍງານບໍ່ມີຂໍ້ມູນ', 'ຍັງບໍ່ Run TE Calculation', 'Run Calculation ແລະກວດ Filter'],
        ['TE ບໍ່ກົງ', 'ສູດ/Benchmark ຕ່າງກັນ', 'ກວດ Row-by-row ແລະປຶກສາ Expert'],
    ]
)

# Section 15 - Good practices
doc.add_paragraph('• ໃຊ້ Template ລ່າສຸດຈາກລະບົບທຸກຄັ້ງ.')
doc.add_paragraph('• ກວດສອບ Province, District, TIN, Tax Year ກ່ອນ Import.')
doc.add_paragraph('• ຫຼັງ Run TE Calculation ຄວນກວດສອບລາຍງານຫຼັກ.')
doc.add_paragraph('• ເມື່ອພົບ Error ໃຫ້ເກັບ Batch ID ແລະ Log ໄວ້.')

# Section 17 - Menu Reference
add_heading('17. ພາກຜະນວກ: ຕໍາແໜ່ງເມນູສໍາຄັນ', 1)
add_table(
    ['ວຽກ', 'ເມນູ'],
    [
        ['Import Profit Tax', 'Get Tax Data > Data Requirement > Profit Tax'],
        ['Import PIT', 'Get Tax Data > Data Requirement > Individual Tax'],
        ['Import Domestic VAT', 'Get Tax Data > Data Requirement > Domestic VAT'],
        ['Import SEZ', 'Get Tax Data > Data Requirement > For SEZ Developers/Investors'],
        ['Import Non-Tax', 'Get Tax Data > Data Requirement > Non-Tax'],
        ['Import ASYCUDA', 'Get Tax Data > Data from ASYCUDA > Import New Data'],
        ['Batch Management', 'Get Tax Data > Data Requirement > Batch Management'],
        ['TE Calculation', 'TE Calculation > [choose tax type]'],
        ['TE Reports', 'TE Reports > [choose report type]'],
    ]
)

# Section 18 - Screenshot Checklist
add_heading('18. Screenshot Checklist', 1)
doc.add_paragraph('Screenshots to add later (placeholder boxes throughout the document):')
screenshots = [
    'Login page',
    'Dashboard and sidebar navigation',
    'Download Template (Domestic VAT import page)',
    'Import success and recent batches',
    'Batch Management page',
    'View batch with Go to TE Calculation button',
    'TE Calculation page (before running)',
    'TE Calculation page (after running)',
    'TE by Tax Type report',
    'Report filters and Export buttons',
    'Import error log example',
    'Excel template example',
]
for i, ss in enumerate(screenshots, 1):
    doc.add_paragraph(f'{i}. {ss}')

doc.add_paragraph('')
p = doc.add_paragraph()
run = p.add_run('— End of Document —')
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.font.italic = True

# Save
output_path = 'D:/Tax-ETS/Documentation/user-manual.docx'
doc.save(output_path)
print(f'Saved to {output_path}')
